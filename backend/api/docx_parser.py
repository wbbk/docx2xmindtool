import docx
import os

def parse_docx(file_path):
    """
    解析docx文件，提取文档结构和内容
    
    Args:
        file_path: docx文件路径
        
    Returns:
        dict: 包含文档结构和内容的字典
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
        
    try:
        # 使用python-docx库打开文档
        doc = docx.Document(file_path)
        
        # 提取文档标题（如果有）
        title = doc.core_properties.title
        if not title and doc.paragraphs:
            # 如果没有设置标题属性，使用第一段作为标题
            title = doc.paragraphs[0].text
            
        # 提取文档内容
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                # 分析段落样式，判断是否为标题
                style_name = para.style.name
                is_heading = style_name.startswith('Heading') if style_name else False
                level = int(style_name.replace('Heading ', '')) if is_heading else 0
                
                content.append({
                    'text': para.text,
                    'is_heading': is_heading,
                    'level': level
                })
        
        # 提取表格内容
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'title': title,
            'content': content,
            'tables': tables
        }
        
    except Exception as e:
        raise Exception(f"解析docx文件失败: {str(e)}")